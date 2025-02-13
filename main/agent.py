
from pydantic_ai.models.openai import OpenAIModel
from docx import Document
from pydantic_ai import Agent, RunContext
import nest_asyncio
import os

nest_asyncio.apply()


model = OpenAIModel('gpt-4o', api_key=os.environ.get('OPEN_AI_KEY'))


agent = Agent(
    model=model,
    system_prompt=(
        'Eres un asistente de un curso, cuya informacion se encuentra en el sylabus del curso'
        'para responder solo puedes utilizar la informacion que te proporciona la funcion get_info'
        'en caso de que te pregunten informacion que no se encuentra en el documento, responde utilizando la funcion skip_message'
    ),
)


@agent.tool
def get_info(ctx: RunContext[str]):
    """
    Lee el contenido de un archivo .docx, incluyendo párrafos y tablas.

    :param ruta_archivo: Ruta al archivo de Word (.docx).
    :return: Contenido completo del archivo como una cadena de texto.
    """
    try:
        # Cargar el documento
        documento = Document('main/syllabus.docx')

        contenido = []

        # Leer los párrafos
        for parrafo in documento.paragraphs:
            if parrafo.text.strip():  # Evitar líneas vacías
                contenido.append(parrafo.text)

        # Leer las tablas
        for tabla in documento.tables:
            for fila in tabla.rows:
                fila_texto = [celda.text.strip() for celda in fila.cells]
                contenido.append("\t".join(fila_texto))  # Separar celdas con tabulación

        # Unir todo el contenido en un solo texto
        return "\n".join(contenido)

    except Exception as e:
        print(f"Error al leer el archivo: {e}")
        return ""


def run_agent(promt:str):
    result_sync = agent.run_sync(
        promt, model_settings={'temperature': 0.0}
        )
    print('respuesta:', result_sync.data)
    print('-'*50)
    return result_sync.data


if __name__ == '__main__':
    while True:
        promt = input('Escribe tu pregunta: ')
        result_sync = agent.run_sync(
            promt, model_settings={'temperature': 0.0}
        )
        print('respuesta:', result_sync.data)
        print('-'*50)